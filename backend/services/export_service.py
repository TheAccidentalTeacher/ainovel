"""
Export service for generating DOCX manuscripts from chapters.

Combines all chapters into a formatted Word document with title page,
table of contents, and proper chapter formatting.
"""

from typing import Optional
from io import BytesIO
import structlog
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from models.schemas import Project, Premise, Chapter

logger = structlog.get_logger()


def generate_manuscript_docx(
    project: Project,
    premise: Premise,
    chapters: list[Chapter],
) -> BytesIO:
    """
    Generate a complete manuscript DOCX file from chapters.
    
    Args:
        project: Project metadata
        premise: Project premise with genre info
        chapters: List of chapters in order
        
    Returns:
        BytesIO buffer containing DOCX file
    """
    logger.info(
        "generating_manuscript_docx",
        project_id=project.id,
        chapter_count=len(chapters),
        total_words=sum(ch.word_count for ch in chapters),
    )
    
    # Create document
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Title Page
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(project.title)
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    
    doc.add_paragraph()  # Spacing
    
    # Author/Genre info
    author_para = doc.add_paragraph()
    author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author_run = author_para.add_run(f"{premise.genre}")
    if premise.subgenre:
        author_run.text += f" / {premise.subgenre}"
    author_run.font.size = Pt(14)
    
    doc.add_paragraph()  # Spacing
    
    # Word count
    total_words = sum(ch.word_count for ch in chapters)
    word_count_para = doc.add_paragraph()
    word_count_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    word_count_run = word_count_para.add_run(f"{total_words:,} words")
    word_count_run.font.size = Pt(12)
    word_count_run.font.italic = True
    
    # Page break before chapters
    doc.add_page_break()
    
    # Add each chapter
    for i, chapter in enumerate(sorted(chapters, key=lambda x: x.chapter_index)):
        logger.debug(
            "adding_chapter_to_docx",
            chapter_index=chapter.chapter_index,
            title=chapter.title,
            word_count=chapter.word_count,
        )
        
        # Chapter heading
        heading = doc.add_heading(level=1)
        heading_run = heading.add_run(f"Chapter {chapter.chapter_index}: {chapter.title}")
        heading_run.font.size = Pt(18)
        heading_run.font.bold = True
        
        # Chapter content
        paragraphs = chapter.content.split('\n\n')
        for para_text in paragraphs:
            if para_text.strip():
                para = doc.add_paragraph(para_text.strip())
                para.paragraph_format.first_line_indent = Inches(0.5)
                para.paragraph_format.space_after = Pt(6)
        
        # Page break after each chapter (except last)
        if i < len(chapters) - 1:
            doc.add_page_break()
    
    # Save to BytesIO buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    logger.info(
        "manuscript_docx_generated",
        project_id=project.id,
        chapters=len(chapters),
        total_words=total_words,
        buffer_size=len(buffer.getvalue()),
    )
    
    return buffer


def generate_manuscript_markdown(
    project: Project,
    premise: Premise,
    chapters: list[Chapter],
) -> str:
    """
    Generate a complete manuscript in Markdown format.
    
    Args:
        project: Project metadata
        premise: Project premise with genre info
        chapters: List of chapters in order
        
    Returns:
        String containing full manuscript in Markdown
    """
    logger.info(
        "generating_manuscript_markdown",
        project_id=project.id,
        chapter_count=len(chapters),
        total_words=sum(ch.word_count for ch in chapters),
    )
    
    # Build markdown content
    lines = []
    
    # Title page
    lines.append(f"# {project.title}\n")
    lines.append(f"**Genre:** {premise.genre}")
    if premise.subgenre:
        lines.append(f" / {premise.subgenre}")
    lines.append("\n")
    
    # Word count
    total_words = sum(ch.word_count for ch in chapters)
    lines.append(f"*{total_words:,} words*\n")
    lines.append("\n---\n\n")
    
    # Add each chapter
    for chapter in sorted(chapters, key=lambda x: x.chapter_index):
        logger.debug(
            "adding_chapter_to_markdown",
            chapter_index=chapter.chapter_index,
            title=chapter.title,
            word_count=chapter.word_count,
        )
        
        # Chapter heading
        lines.append(f"## Chapter {chapter.chapter_index}: {chapter.title}\n\n")
        
        # Chapter content
        paragraphs = chapter.content.split('\n\n')
        for para_text in paragraphs:
            if para_text.strip():
                lines.append(f"{para_text.strip()}\n\n")
        
        # Section break between chapters
        lines.append("\n---\n\n")
    
    markdown_content = "".join(lines)
    
    logger.info(
        "manuscript_markdown_generated",
        project_id=project.id,
        chapters=len(chapters),
        total_words=total_words,
        content_length=len(markdown_content),
    )
    
    return markdown_content
